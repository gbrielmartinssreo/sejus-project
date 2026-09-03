"""Inspecao e preenchimento seguro dos templates DOCX da SEJUS."""
from __future__ import annotations

from pathlib import Path
from uuid import uuid4

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[3]
TEMPLATES_DIR = PROJECT_ROOT / "docs" / "templates"
OUTPUTS_DIR = PROJECT_ROOT / "outputs"


class TemplateError(ValueError):
    """Erro de validacao ao selecionar ou preencher um template."""


def list_templates() -> list[str]:
    """Lista somente os templates DOCX disponiveis no diretorio oficial."""
    if not TEMPLATES_DIR.exists():
        return []
    return sorted(path.name for path in TEMPLATES_DIR.glob("*.docx"))


def resolve_template(template_name: str) -> Path:
    """Resolve um template pelo nome, sem permitir escapar do diretorio."""
    requested = Path(template_name)
    if (
        requested.is_absolute()
        or requested.name != template_name
        or ".." in requested.parts
    ):
        raise TemplateError(f"Template invalido: {template_name}")
    candidate = (TEMPLATES_DIR / requested).resolve()
    base = TEMPLATES_DIR.resolve()
    if candidate.parent != base or candidate.suffix.lower() != ".docx":
        raise TemplateError(f"Template invalido: {template_name}")
    if not candidate.is_file():
        available = ", ".join(list_templates()) or "nenhum"
        raise TemplateError(
            f"Template '{template_name}' nao encontrado. Disponiveis: {available}"
        )
    return candidate


def _paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _placeholder_tokens(document) -> list[str]:
    tokens: set[str] = set()
    for paragraph in _paragraphs(document):
        text = paragraph.text
        starts: list[int] = []
        for index, character in enumerate(text):
            if character == "[":
                starts.append(index)
            elif character == "]" and starts:
                start = starts.pop()
                token = text[start:index + 1]
                if len(token) >= 3 and "\n" not in token:
                    tokens.add(token)
    return sorted(tokens, key=lambda token: (-len(token), token))


def inspect_template(template_name: str) -> dict:
    """Retorna os campos textuais encontrados no corpo e nas tabelas."""
    path = resolve_template(template_name)
    document = Document(str(path))
    placeholders = _placeholder_tokens(document)
    return {
        "template": path.name,
        "placeholders": placeholders,
        "n_placeholders": len(placeholders),
    }


def _replace_in_paragraph(paragraph, values: dict[str, str]) -> int:
    """Substitui tokens entre runs, mantendo a formatacao do run inicial."""
    replacements = 0
    for token, value in values.items():
        if not isinstance(value, str):
            raise TemplateError(f"Valor do campo {token!r} deve ser texto")
        while True:
            runs = paragraph.runs
            combined = "".join(run.text for run in runs)
            match_start = combined.find(token)
            if match_start < 0:
                break
            match_end = match_start + len(token)

            start = end = None
            cumulative_end = 0
            start_offset = end_offset = 0
            for index, run in enumerate(runs):
                next_end = cumulative_end + len(run.text)
                if start is None and match_start < next_end:
                    start = index
                    start_offset = match_start - cumulative_end
                if match_end <= next_end:
                    end = index
                    end_offset = match_end - cumulative_end
                    break
                cumulative_end = next_end

            if start is None or end is None:
                raise TemplateError(f"Nao foi possivel localizar o campo {token}")
            if start == end:
                text = runs[start].text
                runs[start].text = text[:start_offset] + value + text[end_offset:]
            else:
                runs[start].text = runs[start].text[:start_offset] + value
                for index in range(start + 1, end):
                    runs[index].text = ""
                runs[end].text = runs[end].text[end_offset:]
            replacements += 1
    return replacements


def fill_template(template_name: str, values: dict[str, str]) -> dict:
    """Preenche uma copia do template e retorna metadados do arquivo gerado."""
    if not isinstance(values, dict):
        raise TemplateError("values deve ser um objeto com token e valor")

    path = resolve_template(template_name)
    document = Document(str(path))
    placeholders = _placeholder_tokens(document)
    unknown = sorted(set(values) - set(placeholders))
    if unknown:
        raise TemplateError(
            "Campos nao encontrados no template: " + ", ".join(unknown)
        )

    replacements = sum(
        _replace_in_paragraph(paragraph, values)
        for paragraph in _paragraphs(document)
    )
    OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUTS_DIR / f"{path.stem}_{uuid4().hex[:8]}.docx"
    document.save(str(output_path))
    remaining = _placeholder_tokens(Document(str(output_path)))
    return {
        "output_path": str(output_path),
        "template": path.name,
        "replacements": replacements,
        "remaining_placeholders": remaining,
    }