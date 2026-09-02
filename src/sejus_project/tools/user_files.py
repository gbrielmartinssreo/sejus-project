"""
user_files.py
-------------
Tool que lê um arquivo enviado pelo usuário (pasta `importacoes_usuario/`)
e devolve o conteúdo extraído para o agente.

O agente é quem decide, depois de ver o conteúdo, se vai chamar a tool
`consultar_atos_sejus` para comparar com os atos normativos indexados e
apontar o que está faltando na minuta. Este módulo não faz comparação
nem chama nenhum LLM -- só lê e extrai texto.

Formatos suportados: .txt, .md, .pdf, .docx

Uso standalone:
    python user_files.py minuta_contrato.pdf
"""
from __future__ import annotations

import json
from pathlib import Path

# Pasta onde os arquivos enviados pelo usuário ficam. Ajuste se o caminho
# real do projeto for diferente.
IMPORTACOES_DIR = Path("importacoes_usuario")

# Limite de caracteres devolvidos ao agente, para não estourar o contexto
# em arquivos muito grandes. Ajuste conforme necessário.
MAX_CHARS = 20_000

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


class UserFileError(Exception):
    pass


def _list_available_files() -> list[str]:
    if not IMPORTACOES_DIR.exists():
        return []
    return sorted(
        f.name for f in IMPORTACOES_DIR.iterdir()
        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
    )


def _resolve_file(filename: str) -> Path:
    """Resolve o nome do arquivo dentro de importacoes_usuario/, protegendo
    contra path traversal (ex: '../../etc/passwd')."""
    candidate = (IMPORTACOES_DIR / Path(filename).name).resolve()
    base = IMPORTACOES_DIR.resolve()

    if base not in candidate.parents and candidate != base:
        raise UserFileError(f"Caminho inválido: {filename}")

    if not candidate.exists():
        available = _list_available_files()
        hint = f" Arquivos disponíveis: {', '.join(available)}" if available else " Nenhum arquivo encontrado na pasta."
        raise UserFileError(f"Arquivo '{filename}' não encontrado em {IMPORTACOES_DIR}/.{hint}")

    return candidate


def _extract_txt_or_md(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise UserFileError(
            "Leitura de PDF requer a biblioteca 'pypdf'. Instale com: pip install pypdf"
        ) from e

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages)


def _extract_docx(path: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError as e:
        raise UserFileError(
            "Leitura de DOCX requer a biblioteca 'python-docx'. Instale com: pip install python-docx"
        ) from e

    document = docx.Document(str(path))
    paragraphs = [p.text for p in document.paragraphs]

    # inclui texto de tabelas também, já que minutas costumam ter tabelas
    for table in document.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(cell.text for cell in row.cells))

    return "\n".join(paragraphs)


_EXTRACTORS = {
    ".txt": _extract_txt_or_md,
    ".md": _extract_txt_or_md,
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
}


def read_user_file(filename: str) -> dict:
    """Lê e extrai o conteúdo de um arquivo em importacoes_usuario/.

    Devolve um dict com metadados + texto extraído (truncado se muito
    grande), pronto para o agente analisar."""
    path = _resolve_file(filename)
    extension = path.suffix.lower()

    extractor = _EXTRACTORS.get(extension)
    if extractor is None:
        raise UserFileError(
            f"Formato '{extension}' não suportado. Formatos aceitos: "
            f"{', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    text = extractor(path)
    truncated = len(text) > MAX_CHARS
    if truncated:
        text = text[:MAX_CHARS]

    return {
        "filename": path.name,
        "extension": extension,
        "n_chars": len(text),
        "truncated": truncated,
        "text": text,
    }


# ---------------------------------------------------------------------------
# Tool para function calling (formato OpenAI, igual às outras tools do projeto)
# ---------------------------------------------------------------------------

definition = {
    "type": "function",
    "function": {
        "name": "analisar_arquivo_usuario",
        "description": (
            "Lê um arquivo enviado pelo usuário (armazenado na pasta "
            "importacoes_usuario/) e devolve o conteúdo extraído em texto. "
            "Use esta ferramenta quando o usuário pedir para avaliar, revisar "
            "ou verificar uma minuta/documento que ele enviou. Depois de ler "
            "o conteúdo, se for necessário comparar com as normas da SEJUS "
            "(ex: verificar o que está faltando), chame também a ferramenta "
            "consultar_atos_sejus."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": (
                        "Nome do arquivo dentro da pasta importacoes_usuario/ "
                        "(ex: 'minuta_contrato.pdf'). Se não souber o nome "
                        "exato, chame sem preencher para listar os arquivos "
                        "disponíveis."
                    ),
                },
            },
            "required": ["filename"],
        },
    },
}


def analisar_arquivo_usuario(filename: str | None = None) -> str:
    """Função exposta ao agente. Sempre devolve uma string (JSON) --
    nunca lança exceção para o chamador, para o agente conseguir reagir
    ao erro (ex: pedir o nome certo do arquivo) em vez de quebrar."""
    if not filename:
        available = _list_available_files()
        return json.dumps({
            "error": "Nenhum nome de arquivo informado.",
            "available_files": available,
        }, ensure_ascii=False)

    try:
        result = read_user_file(filename)
        return json.dumps(result, ensure_ascii=False)
    except UserFileError as e:
        return json.dumps({"error": str(e)}, ensure_ascii=False)


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Uso: python user_files.py <nome_do_arquivo>")
        print(f"Arquivos disponíveis em {IMPORTACOES_DIR}/: {_list_available_files()}")
        sys.exit(1)

    output = analisar_arquivo_usuario(sys.argv[1])
    parsed = json.loads(output)
    if "error" in parsed:
        print(f"ERRO: {parsed['error']}")
    else:
        print(f"Arquivo: {parsed['filename']} ({parsed['n_chars']} caracteres, truncado={parsed['truncated']})")
        print("---")
        print(parsed["text"][:1000])