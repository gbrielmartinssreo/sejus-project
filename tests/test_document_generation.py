import json
from pathlib import Path

import pytest
from docx import Document

from sejus_project.tools import document_generation as generation
from sejus_project.tools import docx_templates

PROMPTS_DIR = Path(__file__).parent / "prompts"

SCENARIOS = [
    ("portaria_limpeza_cuiaba.txt", "Template_Portaria.docx"),
    ("decreto_higiene_unidades.txt", "Template_Decreto.docx"),
    ("instrucao_rotina_limpeza.txt", "Template_Instrucao_Normativa.docx"),
    ("portaria_conjunta_higiene.txt", "Template_Portaria_Conjunta.docx"),
    ("retificacao_portaria.txt", "Template_Retificacao_Portaria.docx"),
]


@pytest.fixture
def fake_retrieval(monkeypatch):
    calls = []

    def retrieve(query, **kwargs):
        calls.append((query, kwargs))
        return [
            {
                "source_file": "ato_teste.md",
                "act_type": kwargs.get("act_type") or "ATO",
                "act_number": "10/2025",
                "score": 0.91,
                "text": "Fundamento normativo recuperado para a minuta.",
            }
        ]

    monkeypatch.setattr(generation, "retrieve", retrieve)
    return calls


@pytest.mark.parametrize("prompt_name, expected_template", SCENARIOS)
def test_prompt_generates_complete_docx(
    prompt_name, expected_template, fake_retrieval, monkeypatch, tmp_path
):
    request = (PROMPTS_DIR / prompt_name).read_text(encoding="utf-8")
    monkeypatch.setattr(docx_templates, "OUTPUTS_DIR", tmp_path)
    monkeypatch.setattr(generation, "OUTPUTS_DIR", tmp_path, raising=False)

    inspection = json.loads(generation.gerar_documento_normativo(request))

    assert inspection["status"] == "awaiting_confirmation"
    assert inspection["template"] == expected_template
    assert inspection["placeholders"]
    assert fake_retrieval
    assert inspection["context"][0]["source_file"] == "ato_teste.md"

    values = {
        placeholder: f"MINUTA DE TESTE - {index}"
        for index, placeholder in enumerate(inspection["placeholders"], start=1)
    }
    generated = json.loads(
        generation.gerar_documento_normativo(request, values=values)
    )

    assert generated["status"] == "generated"
    output_path = Path(generated["output_path"])
    assert output_path.parent == tmp_path
    assert output_path.is_file()
    assert generated["remaining_placeholders"] == []
    assert generated["replacements"] >= len(values)

    document = Document(str(output_path))
    text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
    )
    assert "MINUTA DE TESTE" in text
    assert "[XX]" not in text


def test_missing_values_do_not_generate_file(fake_retrieval, monkeypatch, tmp_path):
    monkeypatch.setattr(docx_templates, "OUTPUTS_DIR", tmp_path)
    request = (PROMPTS_DIR / "portaria_limpeza_cuiaba.txt").read_text(
        encoding="utf-8"
    )
    inspection = json.loads(generation.gerar_documento_normativo(request))
    one_value = {inspection["placeholders"][0]: "somente um campo"}

    result = json.loads(
        generation.gerar_documento_normativo(request, values=one_value)
    )

    assert result["status"] == "awaiting_confirmation"
    assert result["missing_fields"]
    assert list(tmp_path.glob("*.docx")) == []


def test_authorization_generates_pending_document(
    fake_retrieval, monkeypatch, tmp_path
):
    monkeypatch.setattr(docx_templates, "OUTPUTS_DIR", tmp_path)
    generation._pending_document = None
    request = "Gere uma portaria sobre limpeza da cadeia em Cuiaba."

    first = json.loads(generation.gerar_documento_normativo(request))
    assert first["status"] == "awaiting_confirmation"

    result = json.loads(
        generation.gerar_documento_normativo("pode inventar consultando o banco")
    )

    assert result["status"] == "generated"
    assert result["review_required"] is True
    assert Path(result["output_path"]).is_file()
    assert result["remaining_placeholders"] == []


def test_automatic_draft_uses_formal_text_without_request_command(
    fake_retrieval, monkeypatch, tmp_path
):
    monkeypatch.setattr(docx_templates, "OUTPUTS_DIR", tmp_path)
    generation._pending_document = None
    request = "Gere uma instrucao normativa sobre rotina de limpeza nas unidades."

    json.loads(generation.gerar_documento_normativo(request))
    result = json.loads(
        generation.gerar_documento_normativo("pode inventar consultando o banco")
    )
    document = Document(result["output_path"])
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "Dispõe sobre rotina de limpeza nas unidades." in text
    assert "Dispõe sobre Gere uma" not in text
    assert "CONSIDERANDO Considerando" not in text
    assert "Art. 3º A DEFINIR" not in text
    assert "Secretário de Estado de Justiça" in text


def test_docx_placeholder_can_cross_runs(tmp_path):
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Ementa: [EM")
    paragraph.add_run("ENTA]")
    source = tmp_path / "source.docx"
    document.save(source)

    loaded = Document(source)
    from sejus_project.tools.docx_templates import _replace_in_paragraph

    assert _replace_in_paragraph(loaded.paragraphs[0], {"[EMENTA]": "Limpeza"}) == 1
    assert loaded.paragraphs[0].text == "Ementa: Limpeza"


def test_docx_placeholder_in_table_is_replaced(monkeypatch, tmp_path):
    document = Document()
    cell = document.add_table(rows=1, cols=1).cell(0, 0)
    cell.text = "Responsavel: [NOME]"
    source = tmp_path / "Template_Tabela.docx"
    document.save(source)

    monkeypatch.setattr(docx_templates, "TEMPLATES_DIR", tmp_path)
    monkeypatch.setattr(docx_templates, "OUTPUTS_DIR", tmp_path / "outputs")
    result = docx_templates.fill_template(
        source.name, {"[NOME]": "Equipe de limpeza"}
    )

    generated = Document(result["output_path"])
    assert generated.tables[0].cell(0, 0).text == "Responsavel: Equipe de limpeza"
    assert result["remaining_placeholders"] == []


def test_template_path_traversal_is_rejected():
    with pytest.raises(docx_templates.TemplateError):
        docx_templates.resolve_template("../Template_Portaria.docx")
